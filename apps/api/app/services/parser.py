"""문서 파서 — PDF, DOCX, TXT 지원"""

import io
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


class DocumentParser:
    def parse(self, file_path: str, source_type: str) -> str:
        """파일 경로와 타입을 받아 텍스트를 반환합니다."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"파일을 찾을 수 없습니다: {file_path}")

        if source_type == "pdf":
            return self._parse_pdf(path)
        if source_type == "docx":
            return self._parse_docx(path)
        if source_type == "txt":
            return path.read_text(encoding="utf-8", errors="ignore")

        raise ValueError(f"지원하지 않는 파일 형식: {source_type}")

    def parse_bytes(self, content: bytes, source_type: str) -> str:
        """바이트 데이터를 직접 파싱합니다."""
        if source_type == "pdf":
            reader = PdfReader(io.BytesIO(content))
            return "\n\n".join(
                page.extract_text() or "" for page in reader.pages
            ).strip()

        if source_type == "docx":
            doc = DocxDocument(io.BytesIO(content))
            return "\n\n".join(para.text for para in doc.paragraphs if para.text.strip())

        if source_type == "txt":
            return content.decode("utf-8", errors="ignore").strip()

        raise ValueError(f"지원하지 않는 파일 형식: {source_type}")

    def _parse_pdf(self, path: Path) -> str:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n\n".join(pages).strip()

    def _parse_docx(self, path: Path) -> str:
        doc = DocxDocument(str(path))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)
